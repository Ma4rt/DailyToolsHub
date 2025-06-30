from flask import Blueprint, request, send_file, jsonify, after_this_request, make_response
import os
from io import BytesIO
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import tempfile
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.lib.units import inch
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from PIL import Image
from reportlab.lib.utils import ImageReader
from moviepy.editor import VideoFileClip
import subprocess
import uuid
import threading
import glob

def docx_to_pdf(docx_file):
    document = Document(docx_file)
    pdf_buffer = BytesIO()
    pdf = canvas.Canvas(pdf_buffer, pagesize=letter)
    width, height = letter
    left_margin = 40
    right_margin = 40
    max_width = width - left_margin - right_margin
    y = height - 40
    font_name = "Helvetica"
    code_font = "Courier"
    font_size = 12
    line_height = 20
    para_spacing = 10  # Espaço extra entre parágrafos
    pdf.setFont(font_name, font_size)
    for para in document.paragraphs:
        text = para.text
        if text.strip() == "":
            # Parágrafo vazio: adiciona espaçamento extra
            y -= line_height // 2
            if y < 40:
                pdf.showPage()
                pdf.setFont(font_name, font_size)
                y = height - 40
            continue
        # Detecta bloco de código: começa com espaço/tab OU estilo de parágrafo
        style_name = getattr(getattr(para, 'style', None), 'name', None)
        is_code = (
            text.startswith(" ") or text.startswith("\t") or
            (style_name and isinstance(style_name, str) and style_name.lower() in ["code", "código", "pré-formatado", "preformatted"])
        )
        # Calcula indentação extra baseada no recuo do parágrafo
        left_indent = para.paragraph_format.left_indent.pt if para.paragraph_format.left_indent else 0
        effective_left = left_margin + left_indent
        lines = text.splitlines() if "\n" in text else [text]
        for line in lines:
            if is_code:
                pdf.setFont(code_font, font_size)
                # Preserva todos os espaços/tabs do início da linha
                raw_line = line.replace("\t", "    ")  # Tab = 4 espaços
                while raw_line:
                    for i in range(len(raw_line), 0, -1):
                        if stringWidth(raw_line[:i], code_font, font_size) <= max_width - (effective_left - left_margin):
                            break
                    pdf.drawString(effective_left, y, raw_line[:i])
                    y -= line_height
                    if y < 40:
                        pdf.showPage()
                        pdf.setFont(code_font, font_size)
                        y = height - 40
                    raw_line = raw_line[i:]
            else:
                pdf.setFont(font_name, font_size)
                words = line.split()
                current_line = ""
                for word in words:
                    test_line = current_line + (" " if current_line else "") + word
                    if stringWidth(test_line, font_name, font_size) <= max_width - (effective_left - left_margin):
                        current_line = test_line
                    else:
                        pdf.drawString(effective_left, y, current_line)
                        y -= line_height
                        if y < 40:
                            pdf.showPage()
                            pdf.setFont(font_name, font_size)
                            y = height - 40
                        current_line = word
                if current_line:
                    pdf.drawString(effective_left, y, current_line)
                    y -= line_height
                    if y < 40:
                        pdf.showPage()
                        pdf.setFont(font_name, font_size)
                        y = height - 40
        # Espaço extra entre parágrafos
        y -= para_spacing
        if y < 40:
            pdf.showPage()
            pdf.setFont(font_name, font_size)
            y = height - 40
    pdf.save()
    pdf_buffer.seek(0)
    return pdf_buffer

def pdf_to_docx(pdf_file):
    reader = PdfReader(pdf_file)
    docx_doc = DocxDocument()
    for page in reader.pages:
        text = page.extract_text()
        if text:
            # Quebra o texto em linhas e adiciona cada linha como parágrafo
            for line in text.splitlines():
                if line.strip() == "":
                    docx_doc.add_paragraph("")
                else:
                    docx_doc.add_paragraph(line)
    output = BytesIO()
    docx_doc.save(output)
    output.seek(0)
    return output

def txt_to_pdf(txt_file):
    pdf_buffer = BytesIO()
    pdf = canvas.Canvas(pdf_buffer, pagesize=letter)
    width, height = letter
    left_margin = 40
    right_margin = 40
    max_width = width - left_margin - right_margin
    y = height - 40
    font_name = "Helvetica"
    font_size = 12
    line_height = 20
    pdf.setFont(font_name, font_size)
    for line in txt_file.read().decode("utf-8").splitlines():
        # Quebra automática de linha
        words = line.split()
        current_line = ""
        for word in words:
            test_line = current_line + (" " if current_line else "") + word
            if stringWidth(test_line, font_name, font_size) <= max_width:
                current_line = test_line
            else:
                pdf.drawString(left_margin, y, current_line)
                y -= line_height
                if y < 40:
                    pdf.showPage()
                    pdf.setFont(font_name, font_size)
                    y = height - 40
                current_line = word
        if current_line:
            pdf.drawString(left_margin, y, current_line)
            y -= line_height
            if y < 40:
                pdf.showPage()
                pdf.setFont(font_name, font_size)
                y = height - 40
    pdf.save()
    pdf_buffer.seek(0)
    return pdf_buffer

def images_to_pdf(files):
    pdf_buffer = BytesIO()
    pdf = canvas.Canvas(pdf_buffer, pagesize=letter)
    width, height = letter
    for file in files:
        img = Image.open(file)
        img_width, img_height = img.size
        ratio = min(width / img_width, height / img_height)
        new_width = int(img_width * ratio)
        new_height = int(img_height * ratio)
        x = (width - new_width) // 2
        y = (height - new_height) // 2
        img = img.convert("RGB")
        image_reader = ImageReader(img)
        pdf.drawImage(image_reader, x, y, new_width, new_height)
        pdf.showPage()
    pdf.save()
    pdf_buffer.seek(0)
    return pdf_buffer

def video_to_audio(video_file):
    # Salva o vídeo temporariamente
    with tempfile.NamedTemporaryFile(delete=False, suffix='.mp4') as temp_video:
        temp_video.write(video_file.read())
        temp_video_path = temp_video.name
    # Extrai o áudio
    audio_buffer = BytesIO()
    with VideoFileClip(temp_video_path) as video:
        audio = video.audio
        if audio is None:
            raise ValueError('O vídeo não possui faixa de áudio.')
        # Salva o áudio como MP3 em um arquivo temporário
        with tempfile.NamedTemporaryFile(delete=False, suffix='.mp3') as temp_audio:
            audio.write_audiofile(temp_audio.name)
            temp_audio.seek(0)
            audio_buffer.write(temp_audio.read())
    audio_buffer.seek(0)
    # Remove arquivos temporários
    os.remove(temp_video_path)
    os.remove(temp_audio.name)
    return audio_buffer

def download_video():
    data = request.get_json()
    url = data.get('url')
    tipo = data.get('tipo', 'video')  # 'video' ou 'audio'
    qualidade = data.get('qualidade', 'alta')  # 'alta', 'media', 'baixa'
    if not url:
        return jsonify({'error': 'URL não fornecida'}), 400
    # Define formato yt-dlp
    if tipo == 'audio':
        ytdlp_format = 'bestaudio/best'
        ext = 'mp3'
        extra_args = ['--extract-audio', '--audio-format', 'mp3']
    else:
        if qualidade == 'alta':
            ytdlp_format = 'bestvideo[height<=1080]+bestaudio/best[height<=1080]'
            ext = 'mp4'
        elif qualidade == 'media':
            ytdlp_format = 'bestvideo[height<=720]+bestaudio/best[height<=720]'
            ext = 'mp4'
        else:
            ytdlp_format = 'worstvideo+worstaudio/worst'
            ext = 'mp4'
        extra_args = []
    # Arquivo temporário
    unique_id = uuid.uuid4().hex
    out_name_template = f"video_{unique_id}.%(ext)s"
    out_path_template = os.path.join(tempfile.gettempdir(), out_name_template)
    # Tenta usar 'yt-dlp', se não encontrar, usa 'python -m yt_dlp'
    import shutil
    yt_dlp_cmd = shutil.which('yt-dlp')
    if yt_dlp_cmd:
        cmd = [yt_dlp_cmd, '-f', ytdlp_format, '-o', out_path_template, url] + extra_args
    else:
        cmd = ['python', '-m', 'yt_dlp', '-f', ytdlp_format, '-o', out_path_template, url] + extra_args
    try:
        subprocess.run(cmd, check=True)
        # Após o download, localizar o arquivo real gerado
        search_pattern = os.path.join(tempfile.gettempdir(), f"video_{unique_id}.*")
        files_found = glob.glob(search_pattern)
        if not files_found:
            return jsonify({'error': 'Arquivo de vídeo não encontrado após o download.'}), 500
        out_path_real = files_found[0]  # deve ser só um arquivo
        out_name_real = os.path.basename(out_path_real)
    except subprocess.CalledProcessError as e:
        return jsonify({'error': 'Erro ao baixar o vídeo', 'details': str(e)}), 500
    file_handle = open(out_path_real, 'rb')
    def remove_later(path, delay=5):
        def _remove():
            import time
            time.sleep(delay)
            try:
                os.remove(path)
            except Exception as e:
                print(f"Erro ao remover arquivo temporário: {e}")
        threading.Thread(target=_remove, daemon=True).start()
    @after_this_request
    def remove_file(response):
        try:
            remove_later(out_path_real, delay=5)
        except Exception as e:
            print(f"Erro ao agendar remoção do arquivo temporário: {e}")
        file_size = os.path.getsize(out_path_real)
        response = make_response(send_file(file_handle, as_attachment=True, download_name=out_name_real))
        response.headers['Content-Length'] = str(file_size)
        return response
    return remove_file(None)

def register_routes(app):
    @app.route('/convert/docx-to-pdf', methods=['POST'])
    def convert_docx_to_pdf():
        if 'file' not in request.files:
            return jsonify({'error': 'No file part'}), 400
        file = request.files['file']
        if not file or not file.filename or not file.filename.lower().endswith('.docx'):
            return jsonify({'error': 'File is not a DOCX'}), 400
        pdf_buffer = docx_to_pdf(file)
        return send_file(
            pdf_buffer,
            as_attachment=True,
            download_name=(file.filename.rsplit('.', 1)[0] + '.pdf') if file and file.filename else 'converted.pdf',
            mimetype='application/pdf'
        )

    @app.route('/convert/pdf-to-docx', methods=['POST'])
    def convert_pdf_to_docx():
        if 'file' not in request.files:
            return jsonify({'error': 'No file part'}), 400
        file = request.files['file']
        if not file or not file.filename or not file.filename.lower().endswith('.pdf'):
            return jsonify({'error': 'File is not a PDF'}), 400
        docx_buffer = pdf_to_docx(file)
        return send_file(
            docx_buffer,
            as_attachment=True,
            download_name=(file.filename.rsplit('.', 1)[0] + '.docx') if file and file.filename else 'converted.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    @app.route('/convert/txt-to-pdf', methods=['POST'])
    def convert_txt_to_pdf():
        if 'file' not in request.files:
            return jsonify({'error': 'No file part'}), 400
        file = request.files['file']
        if not file or not file.filename or not file.filename.lower().endswith('.txt'):
            return jsonify({'error': 'File is not a TXT'}), 400
        pdf_buffer = txt_to_pdf(file)
        return send_file(
            pdf_buffer,
            as_attachment=True,
            download_name=(file.filename.rsplit('.', 1)[0] + '.pdf') if file and file.filename else 'converted.pdf',
            mimetype='application/pdf'
        )

    @app.route('/convert/images-to-pdf', methods=['POST'])
    def convert_images_to_pdf():
        if 'files' not in request.files:
            return jsonify({'error': 'No files part'}), 400
        files = request.files.getlist('files')
        if not files or len(files) == 0:
            return jsonify({'error': 'No selected files'}), 400
        pdf_buffer = images_to_pdf(files)
        return send_file(
            pdf_buffer,
            as_attachment=True,
            download_name='imagens_convertidas.pdf',
            mimetype='application/pdf'
        )

    @app.route('/convert/video-to-audio', methods=['POST'])
    def convert_video_to_audio():
        if 'file' not in request.files:
            return jsonify({'error': 'No file part'}), 400
        file = request.files['file']
        if not file or not file.filename or not file.filename.lower().endswith(('.mp4', '.avi', '.mov', '.mkv')):
            return jsonify({'error': 'File is not a supported video format'}), 400
        try:
            audio_buffer = video_to_audio(file)
        except Exception as e:
            return jsonify({'error': str(e)}), 400
        return send_file(
            audio_buffer,
            as_attachment=True,
            download_name=(file.filename.rsplit('.', 1)[0] + '.mp3') if file and file.filename else 'audio_extraido.mp3',
            mimetype='audio/mpeg'
        )

    app.add_url_rule('/download/video', 'download_video', download_video, methods=['POST'])

    # Aqui serão registradas as rotas de conversão de arquivos
    pass 